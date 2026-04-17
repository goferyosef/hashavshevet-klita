import os
from typing import List

import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from models.document import DocumentRecord

REPORTS_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "reports")

EXCEL_HEADERS = [
    "שם קובץ", "תיקייה", "ספק", "סוג מסמך", "תאריך",
    "סכום כולל", 'מע"מ', 'שיעור מע"מ %', "פעולה", "חותמת זמן", "מספר מסמך",
]

ACTION_COLORS = {
    "הועלה": "D4EDDA",   # green
    "דולג":  "FFF3CD",   # yellow
    "שגיאה": "F8D7DA",   # red
}


class Reporter:
    def __init__(self, output_dir: str = REPORTS_DIR):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    # ── Excel ─────────────────────────────────────────────────────────────────

    def generate_excel(self, records: List[DocumentRecord], run_timestamp: str) -> str:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "מסמכים"
        ws.sheet_view.rightToLeft = True

        # Header row
        header_font = Font(bold=True, color="FFFFFF", name="Arial", size=11)
        header_fill = PatternFill("solid", fgColor="1F4E79")
        thin = Side(style="thin", color="AAAAAA")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)

        for col, header in enumerate(EXCEL_HEADERS, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = border
        ws.row_dimensions[1].height = 22

        # Data rows
        for row_idx, rec in enumerate(records, 2):
            values = [
                rec.file_name, rec.folder, rec.supplier, rec.doc_type, rec.date,
                rec.total_amount, rec.vat_amount, rec.vat_rate,
                rec.action, rec.timestamp, rec.doc_number,
            ]
            fill_color = ACTION_COLORS.get(rec.action, "FFFFFF")
            fill = PatternFill("solid", fgColor=fill_color)
            for col_idx, val in enumerate(values, 1):
                cell = ws.cell(row=row_idx, column=col_idx, value=val)
                cell.font = Font(name="Arial", size=10)
                cell.alignment = Alignment(horizontal="right", vertical="center")
                cell.fill = fill
                cell.border = border
            # Format currency columns
            for col_idx in (6, 7):
                ws.cell(row=row_idx, column=col_idx).number_format = '#,##0.00'
            ws.cell(row=row_idx, column=8).number_format = '0.00"%"'

        # Auto column width
        for col in ws.columns:
            max_len = max((len(str(c.value or "")) for c in col), default=10)
            ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 45)

        path = os.path.join(self.output_dir, f"דוח_מסמכים_{run_timestamp}.xlsx")
        wb.save(path)
        return path

    # ── Word ──────────────────────────────────────────────────────────────────

    def generate_word(self, records: List[DocumentRecord], run_timestamp: str) -> str:
        doc = Document()

        # Page RTL
        for section in doc.sections:
            section.right_margin = Inches(1)
            section.left_margin = Inches(1)

        def _rtl_para(text, style=None, bold=False, size=None, color=None):
            p = doc.add_paragraph(style=style)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(text)
            run.bold = bold
            run.font.name = "Arial"
            if size:
                run.font.size = Pt(size)
            if color:
                run.font.color.rgb = RGBColor(*color)
            return p

        # Title
        _rtl_para("קליטה לחשבשבת — דוח עיבוד מסמכים", bold=True, size=18)
        _rtl_para(f"תאריך הרצה: {run_timestamp}", size=11)

        doc.add_paragraph()

        # Summary
        uploaded = sum(1 for r in records if r.action == "הועלה")
        skipped  = sum(1 for r in records if r.action == "דולג")
        errors   = sum(1 for r in records if r.action == "שגיאה")
        _rtl_para(
            f"סה\"כ מסמכים: {len(records)}   |   הועלו: {uploaded}   |   דולגו: {skipped}   |   שגיאות: {errors}",
            bold=True, size=12,
        )

        doc.add_paragraph()
        _rtl_para("פירוט מסמכים", bold=True, size=13)
        doc.add_paragraph()

        icons = {"הועלה": "✔", "דולג": "–", "שגיאה": "✘"}
        err_color = (0xC0, 0x00, 0x00)

        for rec in records:
            icon = icons.get(rec.action, "·")
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            name_run = p.add_run(f"{icon}  {rec.file_name}")
            name_run.bold = True
            name_run.font.name = "Arial"
            name_run.font.size = Pt(10)

            detail = (
                f"  |  ספק: {rec.supplier}  |  סוג: {rec.doc_type}"
                f"  |  תאריך: {rec.date}  |  מספר: {rec.doc_number}"
                f"  |  סכום: ₪{rec.total_amount:,.2f}"
                f"  |  מע\"מ: ₪{rec.vat_amount:,.2f} ({rec.vat_rate}%)"
                f"  |  פעולה: {rec.action}"
            )
            detail_run = p.add_run(detail)
            detail_run.font.name = "Arial"
            detail_run.font.size = Pt(9)

            if rec.error_msg:
                err_run = p.add_run(f"  ⚠ {rec.error_msg}")
                err_run.font.name = "Arial"
                err_run.font.size = Pt(9)
                err_run.font.color.rgb = RGBColor(*err_color)

        path = os.path.join(self.output_dir, f"דוח_עיבוד_{run_timestamp}.docx")
        doc.save(path)
        return path
